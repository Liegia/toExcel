from docx import Document
import pandas as pd

# Ange sökvägen till din DOCX-fil
docx_path = "kursutbud.docx"

# Öppna dokumentet
document = Document(docx_path)

# Kontrollera om dokumentet innehåller tabeller
if document.tables:
    # Antag att vi arbetar med den första tabellen
    table = document.tables[0]
    data = []
    
    # Iterera över raderna i tabellen
    for row in table.rows:
        data.append([cell.text.strip() for cell in row.cells])
    
    # Om den första raden innehåller rubriker
    header = data[0]
    rows = data[1:]
    
    # Skapa en DataFrame
    df = pd.DataFrame(rows, columns=header)
    
    # Spara DataFrame till en Excel-fil
    output_file = "utdata.xlsx"
    df.to_excel(output_file, index=False)
    print(f"Tabellen har sparats till {output_file}")
else:
    print("Inga tabeller hittades i dokumentet.")

